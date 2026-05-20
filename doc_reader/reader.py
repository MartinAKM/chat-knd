import base64
import json
import posixpath
import urllib.request
import zipfile
from pathlib import Path

import pypdf
import docx
from docx.oxml.ns import qn as _qn

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

_IMG_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}

_VISION_PROMPT = (
    "Analise esta imagem de um documento ERP e responda SOMENTE com o que for útil, sem explicações.\n\n"
    "Regras:\n"
    "1. Se houver código(s) de erro (ex: KND-00423, ORA-06512): liste apenas os códigos e o código do programa ERP visível (ex: EPRO15). Exemplo de resposta: KND-00423 EPRO15\n"
    "2. Se algo estiver destacado (seta, marcação, círculo, sublinhado): descreva o elemento indicado em no máximo 8 palavras.\n"
    "3. Se for uma tela genérica sem nenhuma indicação específica: responda exatamente: sem conteúdo relevante\n\n"
    "Máximo 2 linhas. Sem introduções nem explicações."
)

# Pre-compute Clark-notation attribute name used by DOCX image references
_R_EMBED = _qn("r:embed")
_W_P     = _qn("w:p")
_W_T     = _qn("w:t")


def extract_text(path: Path) -> str:
    """Plain text extraction with no vision pass (backward-compatible)."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {suffix}")


def extract_with_images(path: Path, ollama_url: str, model: str) -> str:
    """
    Extract document text with image descriptions injected immediately after
    the page (PDF) or paragraph (DOCX) where each image appears.

    Falls back to plain text when vision is unavailable (empty ollama_url /
    model) or the file type carries no images.
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf_with_images(path, ollama_url, model)
    if suffix == ".docx":
        return _read_docx_with_images(path, ollama_url, model)
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {suffix}")


def extract_images(path: Path) -> list[str]:
    """Extract all embedded images as base64 strings (no ordering guarantee)."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_images(path)
    if suffix == ".docx":
        return _extract_docx_images(path)
    return []


def describe_images(images: list[str], ollama_url: str, model: str) -> str:
    """Describe a list of images via Ollama; returns concatenated descriptions."""
    if not images or not ollama_url or not model:
        return ""
    results = [
        d for b64 in images
        if (d := _describe_one(b64, ollama_url, model))
    ]
    return "\n".join(results)


# ── Internal helpers ───────────────────────────────────────────────────────

def _describe_one(b64: str, ollama_url: str, model: str) -> str:
    """Call Ollama to describe a single image. Returns '' on failure or no content."""
    payload = json.dumps({
        "model": model,
        "messages": [{"role": "user", "content": _VISION_PROMPT, "images": [b64]}],
        "stream": False,
    }).encode()
    req = urllib.request.Request(
        f"{ollama_url}/api/chat",
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=180) as r:
            text = json.loads(r.read())["message"]["content"].strip()
            if text and "sem conteúdo relevante" not in text.lower():
                return text
            return ""
    except Exception:
        return ""


def _read_pdf(path: Path) -> str:
    reader = pypdf.PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def _read_pdf_with_images(path: Path, ollama_url: str, model: str) -> str:
    """Per-page extraction: page text followed by image descriptions."""
    reader = pypdf.PdfReader(path)
    parts: list[str] = []
    seen: set[str] = set()

    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(text)

        for img in page.images:
            if img.name in seen:
                continue
            seen.add(img.name)
            try:
                b64 = base64.b64encode(img.data).decode()
            except Exception:
                continue
            if ollama_url and model:
                desc = _describe_one(b64, ollama_url, model)
                if desc:
                    parts.append(f"[Imagem: {desc}]")

    return "\n".join(parts)


def _read_docx_with_images(path: Path, ollama_url: str, model: str) -> str:
    """
    Per-paragraph extraction: inject each image description immediately after
    the paragraph where the image is anchored (covers body, tables, headers).
    """
    doc = docx.Document(path)

    # Build rId → base64 for all image relationships in the main document part
    img_by_rid: dict[str, str] = {}
    with zipfile.ZipFile(path) as z:
        nameset = set(z.namelist())
        for rel in doc.part.rels.values():
            if not rel.reltype.endswith("/image"):
                continue
            zip_path = posixpath.normpath(f"word/{rel.target_ref}")
            if zip_path not in nameset:
                continue
            if Path(rel.target_ref).suffix.lower() not in _IMG_EXTS:
                continue
            try:
                img_by_rid[rel.rId] = base64.b64encode(z.read(zip_path)).decode()
            except Exception:
                pass

    described_ids: set[str] = set()
    parts: list[str] = []

    # iter(_W_P) visits ALL <w:p> elements in document order, including those
    # inside tables and text boxes, preserving the visual reading sequence.
    #
    # Within each paragraph, we walk direct children in order so that image
    # descriptions land at the exact position where the image appears —
    # not after all the paragraph text as a batch.
    for para_elem in doc._element.iter(_W_P):
        pending: list[str] = []

        for child in para_elem:
            child_text = "".join(e.text or "" for e in child.iter(_W_T)).strip()
            child_rids = [
                elem.get(_R_EMBED) for elem in child.iter()
                if elem.get(_R_EMBED) in img_by_rid
                and elem.get(_R_EMBED) not in described_ids
            ]

            if child_rids:
                # Flush text accumulated before this image, then describe
                if pending:
                    parts.append(" ".join(pending))
                    pending = []
                for rid in child_rids:
                    described_ids.add(rid)
                    if ollama_url and model:
                        desc = _describe_one(img_by_rid[rid], ollama_url, model)
                        if desc:
                            parts.append(f"[Imagem: {desc}]")

            if child_text:
                pending.append(child_text)

        if pending:
            parts.append(" ".join(pending))

    return "\n".join(parts)


def _extract_pdf_images(path: Path) -> list[str]:
    reader = pypdf.PdfReader(path)
    images: list[str] = []
    seen: set[str] = set()
    for page in reader.pages:
        for img in page.images:
            if img.name in seen:
                continue
            seen.add(img.name)
            try:
                images.append(base64.b64encode(img.data).decode())
            except Exception:
                pass
    return images


def _extract_docx_images(path: Path) -> list[str]:
    images: list[str] = []
    try:
        with zipfile.ZipFile(path) as z:
            for name in z.namelist():
                if name.startswith("word/media/") and Path(name).suffix.lower() in _IMG_EXTS:
                    try:
                        images.append(base64.b64encode(z.read(name)).decode())
                    except Exception:
                        pass
    except Exception:
        pass
    return images
